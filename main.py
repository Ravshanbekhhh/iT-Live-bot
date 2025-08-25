import asyncio
import json
import os
from aiogram import Bot, Dispatcher, F
from aiogram.types import Message
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import StatesGroup, State
from docx import Document
from dotenv import load_dotenv

from buttons import admin_btn
from aiogram.types import FSInputFile

load_dotenv()  # .env faylni yuklash
TOKEN = os.getenv("BOT_TOKEN")  #

# --- States ---
class RegisterState(StatesGroup):
    name = State()
    phone = State()

dp = Dispatcher()

ADMINS_FILE = "admins.json"

# --- Adminlarni yuklab olish ---
def load_admins():
    if not os.path.exists(ADMINS_FILE):
        with open(ADMINS_FILE, "w", encoding="utf-8") as f:
            json.dump([7485738561], f)  # asosiy admin id
    with open(ADMINS_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def save_admins(admins):
    with open(ADMINS_FILE, "w", encoding="utf-8") as f:
        json.dump(admins, f, indent=4, ensure_ascii=False)

# --- Admin check ---
def is_admin(user_id: int):
    admins = load_admins()
    return user_id in admins

# /start
@dp.message(Command("start"))
async def command_start_handler(msg: Message, state: FSMContext):
    if is_admin(msg.from_user.id):
        await msg.answer(
            f"Assalomu alaykum {msg.from_user.full_name}! Siz adminsiz ✅",
            reply_markup=admin_btn
        )
    else:
        await state.set_state(RegisterState.name)
        await msg.answer("Ismingizni kiriting:")

# 1-bosqich: ism
@dp.message(RegisterState.name)
async def get_name(message: Message, state: FSMContext):
    await state.update_data(name=message.text)
    await state.set_state(RegisterState.phone)
    await message.answer("📱 Telefon raqamingizni yozing (masalan: +998901234567):")

# 2-bosqich: telefon
@dp.message(RegisterState.phone)
async def phone_handler(message: Message, state: FSMContext):
    await state.update_data(phone=message.text)
    data = await state.get_data()

    users = []
    if os.path.exists("users.json"):
        with open("users.json", "r", encoding="utf-8") as f:
            try:
                users = json.load(f)
            except:
                users = []

    # Telefon raqam bo‘yicha tekshiruv
    for user in users:
        if user["phone"] == data["phone"]:
            await message.answer("❌ Siz allaqachon ro‘yxatdan o‘tgansiz!")
            await state.clear()
            return

    # Yangi foydalanuvchi qo‘shish
    users.append(data)
    with open("users.json", "w", encoding="utf-8") as f:
        json.dump(users, f, indent=4, ensure_ascii=False)

    await message.answer("✅ Ma'lumotlaringiz uchun rahmat! Tez orada siz bilan aloqaga chiqamiz.")
    await state.clear()

# 📋 Obunachilar ro‘yxati
@dp.message(F.text == "Ro'yxatdan otganlar ma'lumotlari")
async def show_users(message: Message):
    if not is_admin(message.from_user.id):
        await message.answer("❌ Bu buyruq faqat adminlar uchun.")
        return

    try:
        with open("users.json", "r", encoding="utf-8") as f:
            users = json.load(f)

        if not users:
            await message.answer("❌ Hozircha hech qanday obunachi yo‘q.")
            return

        # Word fayl
        doc = Document()
        doc.add_heading("📋 Qiziqish bildirganlar", level=1)

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "№"
        hdr_cells[1].text = "Ism"
        hdr_cells[2].text = "Telefon"

        for idx, user in enumerate(users, start=1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = user.get("name", "")
            row_cells[2].text = user.get("phone", "")

        filename = "Leadlar.docx"
        doc.save(filename)

        file = FSInputFile(filename)
        await message.answer_document(file)

        # JSONni bo‘shatamiz
        with open("users.json", "w", encoding="utf-8") as f:
            json.dump([], f)

    except FileNotFoundError:
        await message.answer("❌ Hozircha hech qanday obunachi yo‘q.")

# --- Adminlarni boshqarish ---
@dp.message(Command("addadmin"))
async def add_admin(message: Message):
    if not is_admin(message.from_user.id):
        await message.answer("❌ Sizda huquq yo‘q!")
        return

    parts = message.text.split()
    if len(parts) != 2 or not parts[1].isdigit():
        await message.answer("ℹ️ Foydalanish: /addadmin <user_id>")
        return

    new_admin = int(parts[1])
    admins = load_admins()
    if new_admin in admins:
        await message.answer("⚠️ Bu foydalanuvchi allaqachon admin.")
        return

    admins.append(new_admin)
    save_admins(admins)
    await message.answer(f"✅ {new_admin} admin qilib qo‘shildi.")

@dp.message(Command("deleteadmin"))
async def delete_admin(message: Message):
    if not is_admin(message.from_user.id):
        await message.answer("❌ Sizda huquq yo‘q!")
        return

    parts = message.text.split()
    if len(parts) != 2 or not parts[1].isdigit():
        await message.answer("ℹ️ Foydalanish: /deleteadmin <user_id>")
        return

    admin_id = int(parts[1])
    admins = load_admins()
    if admin_id not in admins:
        await message.answer("❌ Bunday admin topilmadi.")
        return

    if admin_id == 7485738561:
        await message.answer("⚠️ Asosiy adminni o‘chirib bo‘lmaydi.")
        return

    admins.remove(admin_id)
    save_admins(admins)
    await message.answer(f"🗑️ {admin_id} adminlardan o‘chirildi.")

@dp.message(Command("admins"))
async def list_admins(message: Message):
    if not is_admin(message.from_user.id):
        await message.answer("❌ Sizda huquq yo‘q!")
        return

    admins = load_admins()
    text = "👑 Adminlar ro‘yxati:\n\n"
    text += "\n".join([str(a) for a in admins])
    await message.answer(text)

async def main():
    bot = Bot(token=TOKEN)
    # pollingdan oldin webhookni o‘chirib tashlash
    await bot.delete_webhook(drop_pending_updates=True)
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())
